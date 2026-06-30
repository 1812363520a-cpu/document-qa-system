import re
import subprocess
import tempfile
from html import unescape
from io import BytesIO
from pathlib import Path

from docx import Document
from pypdf import PdfReader


class DocumentParseError(ValueError):
    pass


class DocumentParser:
    def parse(self, file_type: str, content: bytes) -> str:
        if file_type == "txt":
            return self._decode_utf8(content)
        if file_type == "markdown":
            text = self._decode_utf8(content)
            return self._markdown_to_plain_text(text)
        if file_type == "pdf":
            return self._pdf_to_plain_text(content)
        if file_type == "doc":
            return self._doc_to_plain_text(content)
        if file_type == "docx":
            return self._docx_to_plain_text(content)
        raise DocumentParseError(f"Unsupported parser document type: {file_type}")

    def _decode_utf8(self, content: bytes) -> str:
        try:
            return content.decode("utf-8")
        except UnicodeDecodeError as exc:
            raise DocumentParseError("Document content must be valid UTF-8 text") from exc

    def _markdown_to_plain_text(self, text: str) -> str:
        without_code_fences = re.sub(r"```.*?```", " ", text, flags=re.DOTALL)
        without_inline_code = re.sub(r"`([^`]*)`", r"\1", without_code_fences)
        without_links = re.sub(r"\[([^\]]+)\]\([^)]+\)", r"\1", without_inline_code)
        without_markup = re.sub(r"^[#>*\-\s]+", "", without_links, flags=re.MULTILINE)
        without_emphasis = without_markup.replace("**", "").replace("__", "")
        without_emphasis = without_emphasis.replace("*", "").replace("_", "")
        return without_emphasis

    def _pdf_to_plain_text(self, content: bytes) -> str:
        try:
            reader = PdfReader(BytesIO(content))
            page_text = [page.extract_text() or "" for page in reader.pages]
        except Exception as exc:
            raise DocumentParseError("PDF content could not be parsed") from exc

        text = "\n".join(page_text).strip()
        if not text:
            raise DocumentParseError("PDF did not contain extractable text")
        return text

    def _doc_to_plain_text(self, content: bytes) -> str:
        if content.startswith(b"PK\x03\x04"):
            return self._docx_to_plain_text(content)

        if content.lstrip().startswith(b"{\\rtf"):
            return self._rtf_to_plain_text(self._decode_document_text(content))

        antiword_detail = ""
        with tempfile.NamedTemporaryFile(suffix=".doc") as temporary_file:
            temporary_file.write(content)
            temporary_file.flush()
            temporary_path = Path(temporary_file.name)

            try:
                result = subprocess.run(
                    ["antiword", str(temporary_path)],
                    capture_output=True,
                    check=False,
                    text=True,
                )
            except FileNotFoundError as exc:
                fallback_text = self._doc_compatibility_text(content)
                if fallback_text:
                    return fallback_text
                raise DocumentParseError(
                    "Word .doc parsing requires the antiword command to be installed"
                ) from exc

        if result.returncode != 0:
            antiword_detail = (result.stderr or result.stdout).strip()
            fallback_text = self._doc_compatibility_text(content)
            if fallback_text:
                return fallback_text
            message = (
                "Word .doc content could not be parsed. Make sure the file is a "
                "legacy Word 97-2003 .doc file, a readable RTF/HTML/text document, "
                "or a renamed .docx file, and is not encrypted or corrupted."
            )
            if antiword_detail:
                message = f"{message} Parser detail: {antiword_detail}"
            raise DocumentParseError(message)

        text = result.stdout.strip()
        if not text:
            raise DocumentParseError("Word .doc document did not contain extractable text")
        return text

    def _doc_compatibility_text(self, content: bytes) -> str:
        try:
            text = self._decode_document_text(content)
        except DocumentParseError:
            return ""

        stripped = text.lstrip()
        if stripped.startswith("{\\rtf"):
            return self._rtf_to_plain_text(text)
        if self._looks_like_html(text):
            return self._html_to_plain_text(text)
        if self._looks_like_text(text):
            return text.strip()
        return ""

    def _decode_document_text(self, content: bytes) -> str:
        encodings = ["utf-8-sig", "gb18030"]
        if content.startswith((b"\xff\xfe", b"\xfe\xff")):
            encodings.insert(0, "utf-16")

        for encoding in encodings:
            try:
                return content.decode(encoding)
            except UnicodeDecodeError:
                continue
        raise DocumentParseError("Document content could not be decoded as text")

    def _looks_like_html(self, text: str) -> bool:
        sample = text[:4096].lower()
        return any(marker in sample for marker in ("<html", "<body", "<!doctype", "<meta"))

    def _looks_like_text(self, text: str) -> bool:
        sample = text[:4096]
        if not sample.strip():
            return False
        control_chars = sum(
            1 for char in sample if ord(char) < 32 and char not in "\n\r\t"
        )
        return control_chars / max(len(sample), 1) < 0.05

    def _html_to_plain_text(self, text: str) -> str:
        without_scripts = re.sub(
            r"<(script|style).*?</\1>", " ", text, flags=re.DOTALL | re.IGNORECASE
        )
        with_line_breaks = re.sub(
            r"</(p|div|br|tr|li|h[1-6])\s*>",
            "\n",
            without_scripts,
            flags=re.IGNORECASE,
        )
        without_tags = re.sub(r"<[^>]+>", " ", with_line_breaks)
        return self._normalize_text(unescape(without_tags))

    def _rtf_to_plain_text(self, text: str) -> str:
        decoded_unicode = re.sub(
            r"\\u(-?\d+)\??",
            lambda match: chr(int(match.group(1)) % 65536),
            text,
        )
        without_hex = re.sub(
            r"\\'[0-9a-fA-F]{2}",
            " ",
            decoded_unicode,
        )
        with_line_breaks = re.sub(r"\\(par|line)\b", "\n", without_hex)
        without_controls = re.sub(r"\\[a-zA-Z]+-?\d* ?", " ", with_line_breaks)
        without_braces = without_controls.replace("{", " ").replace("}", " ")
        return self._normalize_text(without_braces)

    def _normalize_text(self, text: str) -> str:
        lines = [" ".join(line.split()) for line in text.splitlines()]
        return "\n".join(line for line in lines if line).strip()

    def _docx_to_plain_text(self, content: bytes) -> str:
        try:
            document = Document(BytesIO(content))
        except Exception as exc:
            raise DocumentParseError("Word document content could not be parsed") from exc

        parts = []
        parts.extend(paragraph.text for paragraph in document.paragraphs if paragraph.text)
        for table in document.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if cells:
                    parts.append(" | ".join(cells))

        text = "\n".join(parts).strip()
        if not text:
            raise DocumentParseError("Word document did not contain extractable text")
        return text
